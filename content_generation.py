from flask import Flask, request, jsonify, render_template, send_file
import openai
import os
import zipfile
import io
from fpdf import FPDF
from docx import Document
from docx.shared import Inches
from pathlib import Path
from dotenv import load_dotenv

app = Flask(__name__)

load_dotenv()

openai.api_key=os.getenv("OPENAI_API_KEY")


# Function to fetch response from GPT
def fetch_gpt_response(query):
    try:
        response = openai.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are an expert in the pharmaceutical and medical domain only. Only answer those questions and don't answer any other questions."},
                {"role": "user", "content": query},
            ],
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"Error: {str(e)}"

def save_as_pdf(content, file_name="response.pdf", logo_path='assets/logo.jpeg'):
    pdf = FPDF()
    pdf.add_page()

    # Add the logo
    pdf.image(logo_path, x=10, y=8, w=30)

    # Title of the document
    pdf.set_font("Arial", style='B', size=16)
    pdf.ln(30)
    pdf.cell(200, 10, txt="Research Content Response", ln=True, align='C')
    pdf.ln(10)

    # Add content
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(190, 10, content)

    # Save the PDF
    pdf.output(file_name)


def save_as_word(content, file_name="response.docx", logo_path='assets/logo.jpeg'):
    document = Document()

    # Add the logo
    document.add_picture(logo_path, width=1500000)  # Width in EMU (e.g., 1500000 = 150px)

    # Add the title
    document.add_heading("Research Content Response", level=1)

    # Add content
    document.add_paragraph(content)

    # Save the Word document
    document.save(file_name)


def create_scorm_package(content, file_type):
    output_folder = "scorm_package"
    scorm_zip_name = f"scorm_package_{file_type}.zip"

    # Ensure output folder exists
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    # Clear previous files in the folder
    for file in os.listdir(output_folder):
        os.remove(os.path.join(output_folder, file))

    # Save content to the appropriate format
    if file_type == "pdf":
        file_path = os.path.join(output_folder, "content.pdf")
        save_as_pdf(content, file_path)
    elif file_type == "docx":
        file_path = os.path.join(output_folder, "content.docx")
        save_as_word(content, file_path)

    # Add the SCORM-specific files
    imanifest_path = os.path.join(output_folder, "imanifest.xml")
    html_path = os.path.join(output_folder, "index.html")

    with open(imanifest_path, "w") as imanifest_file:
        imanifest_file.write("<manifest>SCORM manifest file content</manifest>")

    with open(html_path, "w") as html_file:
        html_file.write("<html><body><h1>SCORM Content</h1></body></html>")

    # Create SCORM ZIP package
    with zipfile.ZipFile(scorm_zip_name, 'w', zipfile.ZIP_DEFLATED) as scorm_zip:
        scorm_zip.write(file_path, arcname=os.path.basename(file_path))
        scorm_zip.write(imanifest_path, arcname="imanifest.xml")
        scorm_zip.write(html_path, arcname="index.html")

    # Read and return the ZIP file
    with open(scorm_zip_name, 'rb') as scorm_file:
        return scorm_file.read()




# Route to serve the HTML page
@app.route('/')
def index():
    return render_template('index.html')

@app.route('/generate-content', methods=['POST'])
def generate_content():
    try:
        # Get the query from the form
        query = request.form.get("query", "").strip()

        if not query:
            return render_template('index.html', error="Query is required.")

        # Process the query for content generation
        gpt_response = fetch_gpt_response(query)

        return render_template('index.html', query=query, response=gpt_response)

    except Exception as e:
        return render_template('index.html', error=f"Internal server error: {str(e)}")



@app.route('/download-content', methods=['POST'])
def download_content():
    try:
        response = request.form.get("response")
        format_type = request.form.get("format")

        if format_type == "pdf_scorm":
            scorm_file = create_scorm_package(response, "pdf")
            return send_file(io.BytesIO(scorm_file), as_attachment=True, download_name="scorm_package_pdf.zip")
        elif format_type == "docx_scorm":
            scorm_file = create_scorm_package(response, "docx")
            return send_file(io.BytesIO(scorm_file), as_attachment=True, download_name="scorm_package_doc.zip")
        else:
            return render_template('index.html', error="Invalid download format selected.")
    except Exception as e:
        return render_template('index.html', error=f"Error in downloading content: {str(e)}")



if __name__ == '__main__':
    app.run(debug=True)
